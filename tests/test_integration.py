import os
import shutil
import tempfile
import unittest
from typing import List, Dict, Optional
from PIL import Image, UnidentifiedImageError
from docx import Document
from pypdf import PdfReader, PdfWriter
from mutagen.mp3 import MP3
from metadata_cleaner.remover import remove_metadata_from_folder, remove_metadata, dry_run_metadata_removal
from metadata_cleaner.file_handlers.metadata_extractor import extract_metadata
from metadata_cleaner.config.settings import ENABLE_PARALLEL_PROCESSING

class TestMetadataCleaner(unittest.TestCase):
    def setUp(self) -> None:
        """
        Create a temporary test directory with valid test files.
        """
        self.test_dir = tempfile.TemporaryDirectory()
        self.test_folder = self.test_dir.name
        self.test_output_folder = os.path.join(self.test_folder, "output")
        os.makedirs(self.test_output_folder, exist_ok=True)

        # Create test files
        self.test_image = os.path.join(self.test_folder, "test_image.jpg")
        img = Image.new("RGB", (100, 100), color="red")
        img.save(self.test_image, "JPEG")

        self.test_pdf = os.path.join(self.test_folder, "test_document.pdf")
        writer = PdfWriter()
        writer.add_metadata({"/Author": "Test Author"})
        with open(self.test_pdf, "wb") as f:
            writer.write(f)

        self.test_docx = os.path.join(self.test_folder, "test_document.docx")
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.core_properties.author = "Test Author"
        doc.save(self.test_docx)

        self.test_audio = os.path.join(self.test_folder, "test_audio.mp3")
        if shutil.which("ffmpeg"):
            os.system(f"ffmpeg -f lavfi -i anullsrc=r=44100:cl=mono -t 3 -q:a 9 -acodec libmp3lame {self.test_audio} -y")

        self.test_video = os.path.join(self.test_folder, "test_video.mp4")
        if shutil.which("ffmpeg"):
            os.system(f"ffmpeg -f lavfi -i color=c=blue:s=320x240:d=3 -vf format=yuv420p {self.test_video} -y")

    def test_batch_processing(self) -> None:
        """
        Test batch metadata removal for all supported file types.
        """
        cleaned_files: List[str] = remove_metadata_from_folder(self.test_folder, self.test_output_folder)
        expected_count = len([f for f in os.listdir(self.test_folder) if os.path.isfile(os.path.join(self.test_folder, f))])

        self.assertEqual(len(cleaned_files), expected_count, "Not all files were processed correctly.")

        for cleaned_file in cleaned_files:
            self.assertTrue(os.path.exists(cleaned_file), f"Cleaned file does not exist: {cleaned_file}")
            ext = os.path.splitext(cleaned_file)[1].lower()

            if ext in {".jpg", ".jpeg"}:
                try:
                    img = Image.open(cleaned_file)
                    self.assertNotIn("exif", img.info, "EXIF metadata not removed from image.")
                except UnidentifiedImageError:
                    self.fail("Failed to open cleaned image file.")

            elif ext == ".pdf":
                reader = PdfReader(cleaned_file)
                self.assertEqual(reader.metadata, {}, "Metadata not fully removed from PDF.")

            elif ext == ".docx":
                doc = Document(cleaned_file)
                self.assertFalse(doc.core_properties.author, "Metadata not removed from DOCX.")

            elif ext == ".mp3" and shutil.which("ffmpeg"):
                audio = MP3(cleaned_file)
                self.assertFalse(audio.tags, "Metadata not removed from MP3.")

    def test_dry_run_mode(self) -> None:
        """Test dry-run mode to ensure metadata is not actually removed."""
        metadata_before = extract_metadata(self.test_image)
        metadata_after = dry_run_metadata_removal(self.test_image)
        
        self.assertIn("before_removal", metadata_after, "Dry-run output should contain before_removal.")
        self.assertIn("after_removal", metadata_after, "Dry-run output should contain after_removal.")
        self.assertEqual(metadata_before, metadata_after["before_removal"], "Metadata before removal should match original.")
        self.assertNotEqual(metadata_after["before_removal"], metadata_after["after_removal"], "Dry-run should simulate removal.")

    def test_parallel_processing_flag(self) -> None:
        """Test that parallel processing does not exceed system limits."""
        self.assertLessEqual(ENABLE_PARALLEL_PROCESSING, os.cpu_count(), "Parallel workers exceed CPU count.")

    def test_empty_folder(self) -> None:
        """Test behavior when processing an empty folder."""
        empty_folder = os.path.join(self.test_folder, "empty_folder")
        os.makedirs(empty_folder, exist_ok=True)
        cleaned_files = remove_metadata_from_folder(empty_folder, self.test_output_folder)
        self.assertEqual(len(cleaned_files), 0, "Expected zero files to be processed in an empty folder.")

    def tearDown(self) -> None:
        """
        Clean up temporary test directories.
        """
        self.test_dir.cleanup()

if __name__ == "__main__":
    unittest.main()
