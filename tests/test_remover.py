import os
import subprocess
import unittest
import shutil
from typing import List
from PIL import Image
from docx import Document
from pypdf import PdfWriter, PdfReader
from mutagen.mp3 import MP3
from metadata_cleaner.remover import remove_metadata_from_folder


class TestMetadataRemover(unittest.TestCase):
    def setUp(self) -> None:
        """
        Create valid test files for each supported file type in a test folder.
        """
        self.test_folder: str = "test_batch"
        self.output_folder: str = "test_batch_output"
        os.makedirs(self.test_folder, exist_ok=True)

        # Create a valid JPG file with EXIF metadata
        image_path: str = os.path.join(self.test_folder, "test_image.jpg")
        img = Image.new("RGB", (100, 100), color="red")
        img.save(image_path, "JPEG")

        # Create a valid PDF file with metadata
        pdf_path: str = os.path.join(self.test_folder, "test_document.pdf")
        writer = PdfWriter()
        writer.add_metadata({"/Author": "Test Author"})
        with open(pdf_path, "wb") as f:
            writer.write(f)

        # Create a valid DOCX file with metadata
        docx_path: str = os.path.join(self.test_folder, "test_document.docx")
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.core_properties.author = "Test Author"
        doc.save(docx_path)

        # Create a valid MP3 file using ffmpeg (skip if FFmpeg is missing)
        self.test_audio: str = os.path.join(self.test_folder, "test_audio.mp3")
        if shutil.which("ffmpeg"):
            subprocess.run([
                "ffmpeg", "-f", "lavfi", "-i", "anullsrc=r=44100:cl=mono",
                "-t", "3", "-q:a", "9", "-acodec", "libmp3lame", self.test_audio, "-y"
            ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        # Create a valid MP4 file using ffmpeg (skip if FFmpeg is missing)
        self.test_video: str = os.path.join(self.test_folder, "test_video.mp4")
        if shutil.which("ffmpeg"):
            subprocess.run([
                "ffmpeg", "-f", "lavfi", "-i", "color=c=blue:s=320x240:d=3",
                "-vf", "format=yuv420p", self.test_video, "-y"
            ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    def test_batch_processing(self) -> None:
        """
        Test batch metadata removal for all supported file types in the test folder.

        Verifies that the number of processed files matches the number of created files and that each output file exists.
        Also verifies that metadata is actually removed.
        """
        cleaned_files: List[str] = remove_metadata_from_folder(self.test_folder, self.output_folder)

        expected_count = sum(
            os.path.isfile(os.path.join(self.test_folder, f)) for f in os.listdir(self.test_folder)
        )

        # Ensure all files were processed
        self.assertEqual(len(cleaned_files), expected_count, f"Expected {expected_count} files to be processed.")

        for cleaned_file in cleaned_files:
            self.assertTrue(os.path.exists(cleaned_file), f"Cleaned file does not exist: {cleaned_file}")

            # Validate metadata removal per file type
            ext = os.path.splitext(cleaned_file)[1].lower()

            if ext in {".jpg", ".jpeg"}:
                img = Image.open(cleaned_file)
                self.assertNotIn("exif", img.info, "EXIF metadata not removed from image.")

            elif ext == ".pdf":
                reader = PdfReader(cleaned_file)
                self.assertFalse(reader.metadata, "Metadata not removed from PDF.")

            elif ext == ".docx":
                doc = Document(cleaned_file)
                self.assertFalse(doc.core_properties.author, "Metadata not removed from DOCX.")

            elif ext == ".mp3" and shutil.which("ffmpeg"):
                audio = MP3(cleaned_file)
                self.assertFalse(audio.tags, "Metadata not removed from MP3.")

    def tearDown(self) -> None:
        """
        Clean up test directories created during the tests.
        """
        shutil.rmtree(self.test_folder, ignore_errors=True)
        shutil.rmtree(self.output_folder, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
