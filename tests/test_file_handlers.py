import os
import shutil
import tempfile
import unittest
from PIL import Image, UnidentifiedImageError
from docx import Document
from pypdf import PdfReader, PdfWriter
from mutagen.mp3 import MP3
from metadata_cleaner.file_handlers.image.image_handler import remove_metadata
from metadata_cleaner.file_handlers.document.pdf_handler import remove_metadata
from metadata_cleaner.file_handlers.document.docx_handler import remove_metadata
from metadata_cleaner.file_handlers.audio.audio_handler import remove_metadata
from metadata_cleaner.file_handlers.video.video_handler import remove_metadata

class TestFileHandlers(unittest.TestCase):
    def setUp(self):
        """
        Create temporary test files for each supported file type.
        """
        self.test_dir = tempfile.TemporaryDirectory()
        self.test_folder = self.test_dir.name
        self.test_output_folder = os.path.join(self.test_folder, "output")
        os.makedirs(self.test_output_folder, exist_ok=True)

        # Image Test File
        self.test_image = os.path.join(self.test_folder, "test_image.jpg")
        img = Image.new("RGB", (100, 100), color="blue")
        img.save(self.test_image, "JPEG")

        # Corrupt Image File
        self.test_corrupt_image = os.path.join(self.test_folder, "corrupt_image.jpg")
        with open(self.test_corrupt_image, "wb") as f:
            f.write(b"corruptdata")

        # PDF Test File
        self.test_pdf = os.path.join(self.test_folder, "test_document.pdf")
        writer = PdfWriter()
        writer.add_metadata({"/Author": "Test Author"})
        with open(self.test_pdf, "wb") as f:
            writer.write(f)

        # DOCX Test File
        self.test_docx = os.path.join(self.test_folder, "test_document.docx")
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.core_properties.author = "Test Author"
        doc.save(self.test_docx)

        # MP3 Test File
        self.test_audio = os.path.join(self.test_folder, "test_audio.mp3")
        if shutil.which("ffmpeg"):
            os.system(f"ffmpeg -f lavfi -i anullsrc=r=44100:cl=mono -t 3 -q:a 9 -acodec libmp3lame {self.test_audio} -y")

        # MP4 Test File
        self.test_video = os.path.join(self.test_folder, "test_video.mp4")
        if shutil.which("ffmpeg"):
            os.system(f"ffmpeg -f lavfi -i color=c=blue:s=320x240:d=3 -vf format=yuv420p {self.test_video} -y")

    def test_image_handler(self):
        """Test image metadata removal and ensure original remains unchanged."""
        output_file = os.path.join(self.test_output_folder, "test_image_cleaned.jpg")
        result = remove_metadata(self.test_image, output_file)
        self.assertIsNotNone(result)
        self.assertTrue(os.path.exists(result))
        
        img = Image.open(result)
        self.assertNotIn("exif", img.info, "EXIF metadata not removed from image.")

    def test_corrupt_image_handler(self):
        """Test handling of a corrupt image file."""
        output_file = os.path.join(self.test_output_folder, "corrupt_image_cleaned.jpg")
        result = remove_metadata(self.test_corrupt_image, output_file)
        self.assertIsNone(result, "Corrupt image should not be processed.")

    def test_pdf_handler(self):
        """Test PDF metadata removal."""
        output_file = os.path.join(self.test_output_folder, "test_pdf_cleaned.pdf")
        result = remove_metadata(self.test_pdf, output_file)
        self.assertIsNotNone(result)
        self.assertTrue(os.path.exists(result))
        
        reader = PdfReader(result)
        self.assertEqual(reader.metadata, {}, "Metadata not fully removed from PDF.")

    def test_docx_handler(self):
        """Test DOCX metadata removal."""
        output_file = os.path.join(self.test_output_folder, "test_docx_cleaned.docx")
        result = remove_metadata(self.test_docx, output_file)
        self.assertIsNotNone(result)
        self.assertTrue(os.path.exists(result))
        
        doc = Document(result)
        self.assertFalse(doc.core_properties.author, "Metadata not removed from DOCX.")

    def test_audio_handler(self):
        """Test audio metadata removal."""
        if not shutil.which("ffmpeg"):
            self.skipTest("FFmpeg not found, skipping audio test.")
        output_file = os.path.join(self.test_output_folder, "test_audio_cleaned.mp3")
        result = remove_metadata(self.test_audio, output_file)
        self.assertIsNotNone(result)
        self.assertTrue(os.path.exists(result))
        
        audio = MP3(result)
        self.assertFalse(audio.tags, "Metadata not removed from MP3.")

    def test_video_handler(self):
        """Test video metadata removal."""
        if not shutil.which("ffmpeg"):
            self.skipTest("FFmpeg not found, skipping video test.")
        output_file = os.path.join(self.test_output_folder, "test_video_cleaned.mp4")
        result = remove_metadata(self.test_video, output_file)
        self.assertIsNotNone(result)
        self.assertTrue(os.path.exists(result))

    def tearDown(self):
        """
        Clean up temporary test files.
        """
        self.test_dir.cleanup()

if __name__ == "__main__":
    unittest.main()
